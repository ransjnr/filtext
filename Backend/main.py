import os
import mimetypes
import re
import PyPDF2
import pandas as pd
import spacy
import geonamescache
from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from cleaner import clean_data

app = FastAPI()

nlp = spacy.load("en_core_web_sm")

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173"],  # Allow your frontend origin
    allow_credentials=True,
    allow_methods=["*"],  # Allow all HTTP methods
    allow_headers=["*"],  # Allow all headers
)

UPLOAD_DIRECTORY = "uploads"

if not os.path.exists(UPLOAD_DIRECTORY):
    os.makedirs(UPLOAD_DIRECTORY)
    
    # function for preprocessing text
def preprocess_text(text):
# Remove escape characters
    text = re.sub(r'\\[nrtbf]', '', text)  # Remove common escape sequences

# Process the text with spaCy
    doc = nlp(text)

# Extract cleaned tokens (excluding punctuation and stop words)
    cleaned_tokens = [token.text for token in doc if not token.is_punct and not token.is_stop]

# Join the tokens back into a single string
    cleaned_text = ' '.join(cleaned_tokens)
    return cleaned_text

def gen_dict_extract(var, key):
    if isinstance(var, dict):
        for k, v in var.items():
            if k == key:
                yield v
            if isinstance(v, (dict, list)):
                yield from gen_dict_extract(v, key)
    elif isinstance(var, list):
        for d in var:
            yield from gen_dict_extract(d, key)

def spacy_ner_extraction(text):
    global df_sc_name, df_sc_org, df_sc_date, df_sc_city, df_sc_state, df_sc_country, df_sc_time, df_sc_quantity, df_sc_ordinal, df_sc_cardinal, df_sc_percentage, df_sc_event, df_sc_language, df_sc_artwork, df_sc_money, df_sc_loc

    gc = geonamescache.GeonamesCache()
    countries = gc.get_countries()
    cities = gc.get_cities()
    cities = [*gen_dict_extract(cities, 'name')]
    countries = [*gen_dict_extract(countries, 'name')]

    sc_name = []
    sc_org = []
    sc_date = []
    sc_time = []
    sc_city = []
    sc_state = []
    sc_country = []
    sc_loc = []
    sc_artwork = []
    sc_language = []
    sc_money = []
    sc_quantity = []
    sc_ordinal = []
    sc_cardinal = []
    sc_percentage = []
    sc_event = []

    doc = nlp(text)
    for ent in doc.ents:
        if ent.label_ == 'PERSON':
            sc_name.append(ent.text)
        elif ent.label_ == 'ORG':
            sc_org.append(ent.text)
        elif ent.label_ == 'DATE':
            sc_date.append(ent.text)
        elif ent.label_ == 'LOC':
            sc_loc.append(ent.text)
        elif ent.label_ == 'WORK_OF_ART':
            sc_artwork.append(ent.text)
        elif ent.label_ == 'LANGUAGE':
            sc_language.append(ent.text)
        elif ent.label_ == 'PERCENT':
            sc_percentage.append(ent.text)
        elif ent.label_ == 'TIME':
            sc_time.append(ent.text)
        elif ent.label_ == 'MONEY':
            sc_money.append(ent.text)
        elif ent.label_ == 'QUANTITY':
            sc_quantity.append(ent.text)
        elif ent.label_ == 'ORDINAL':
            sc_ordinal.append(ent.text)
        elif ent.label_ == 'CARDINAL':
            sc_cardinal.append(ent.text)
        elif ent.label_ == 'GPE':
            if ent.text in countries:
                sc_country.append(ent.text)
            elif ent.text in cities:
                sc_city.append(ent.text)
            else:
                sc_state.append(ent.text)

    df_sc_name = pd.DataFrame({"Name": sc_name})
    df_sc_org = pd.DataFrame({"Organization": sc_org})
    df_sc_date = pd.DataFrame({"Date": sc_date})
    df_sc_city = pd.DataFrame({"City": sc_city})
    df_sc_state = pd.DataFrame({"State": sc_state})
    df_sc_country = pd.DataFrame({"Country": sc_country})
    df_sc_time = pd.DataFrame({"Time": sc_time})
    df_sc_quantity = pd.DataFrame({"Quantity": sc_quantity})
    df_sc_ordinal = pd.DataFrame({"Ordinal": sc_ordinal})
    df_sc_cardinal = pd.DataFrame({"Cardinal": sc_cardinal})
    df_sc_percentage = pd.DataFrame({"Percentage": sc_percentage})
    df_sc_event = pd.DataFrame({"Event": sc_event})
    df_sc_language = pd.DataFrame({"Language": sc_language})
    df_sc_artwork = pd.DataFrame({"Art Work": sc_artwork})
    df_sc_money = pd.DataFrame({"Money": sc_money})
    df_sc_loc = pd.DataFrame({"Location": sc_loc})

def get_email_addresses(text):
    global df_get_email

    r = re.compile(r'[\w\.-]+@[\w\.-]+')
    email_ori = r.findall(text)
    email = [email.lower().strip(".") for email in email_ori]
    df_get_email = pd.DataFrame({"Email": email})

def read_pdf(file_path_or_stream):
    pdf_text = ""
    reader = PyPDF2.PdfReader(file_path_or_stream)
    for page in reader.pages:
        pdf_text += page.extract_text()
    return pdf_text

@app.post("/upload/")
async def upload_file(file: UploadFile = File(...)):
    try:
        file_path = os.path.join(UPLOAD_DIRECTORY, file.filename)
        with open(file_path, "wb") as buffer:
            buffer.write(await file.read())
        return {"message": "File uploaded successfully"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/extract/")
async def extract_info():
    try:
        file_list = os.listdir(UPLOAD_DIRECTORY)
        if not file_list:
            raise HTTPException(status_code=400, detail="No file uploaded")
        
        file_path = os.path.join(UPLOAD_DIRECTORY, file_list[0])

        mime_type, _ = mimetypes.guess_type(file_path)
        if mime_type == 'application/pdf':
            with open(file_path, "rb") as file:
                text = read_pdf(file)
        elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            from docx import Document
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
        elif mime_type == 'text/plain':
            with open(file_path, "r", encoding='utf-8', errors='replace') as file:
                text = file.read()
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type")
        
        text = preprocess_text(text)
        
        spacy_ner_extraction(text)
        get_email_addresses(text)

        df_output = pd.concat([
            df_sc_name, df_get_email, df_sc_org, df_sc_date, df_sc_city, df_sc_state, df_sc_country, df_sc_percentage,
            df_sc_time, df_sc_quantity, df_sc_cardinal, df_sc_event, df_sc_language, df_sc_artwork, df_sc_money, df_sc_loc
        ], axis=1)

        clean_data_dict = clean_data(df_output.to_dict(orient='records'))
        # Return the extracted text along with cleaned data
        return (df_sc_name)
        return clean_data_dict
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))










# from fastapi import FastAPI, File, UploadFile, HTTPException
# import os
# import mimetypes
# from PyPDF2 import PdfFileReader
# import pandas as pd
# import spacy
# import geonamescache
# import re
# from io import BytesIO
# from fastapi.middleware.cors import CORSMiddleware
# from cleaner import clean_data

# app = FastAPI()

# nlp = spacy.load("en_core_web_sm")
# # Add CORS middleware
# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=["http://localhost:5173"],  # Allow your frontend origin
#     allow_credentials=True,
#     allow_methods=["*"],  # Allow all HTTP methods
#     allow_headers=["*"],  # Allow all headers
# )
# UPLOAD_DIRECTORY = "uploads"

# if not os.path.exists(UPLOAD_DIRECTORY):
#     os.makedirs(UPLOAD_DIRECTORY)

# def gen_dict_extract(var, key):
#     if isinstance(var, dict):
#         for k, v in var.items():
#             if k == key:
#                 yield v
#             if isinstance(v, (dict, list)):
#                 yield from gen_dict_extract(v, key)
#     elif isinstance(var, list):
#         for d in var:
#             yield from gen_dict_extract(d, key)

# def spacy_ner_extraction(text):
#     global df_sc_name, df_sc_org, df_sc_date, df_sc_city, df_sc_state, df_sc_country, df_sc_time, df_sc_quantity, df_sc_ordinal, df_sc_cardinal, df_sc_percentage, df_sc_event, df_sc_language, df_sc_artwork, df_sc_money, df_sc_loc

#     gc = geonamescache.GeonamesCache()
#     countries = gc.get_countries()
#     cities = gc.get_cities()
#     cities = [*gen_dict_extract(cities, 'name')]
#     countries = [*gen_dict_extract(countries, 'name')]

#     sc_name = []
#     sc_org = []
#     sc_date = []
#     sc_time = []
#     sc_city = []
#     sc_state = []
#     sc_country = []
#     sc_loc = []
#     sc_artwork = []
#     sc_language = []
#     sc_money = []
#     sc_quantity = []
#     sc_ordinal = []
#     sc_cardinal = []
#     sc_percentage = []
#     sc_event = []

#     doc = nlp(text)
#     for ent in doc.ents:
#         if ent.label_ == 'PERSON':
#             sc_name.append(ent.text)
#         elif ent.label_ == 'ORG':
#             sc_org.append(ent.text)
#         elif ent.label_ == 'DATE':
#             sc_date.append(ent.text)
#         elif ent.label_ == 'LOC':
#             sc_loc.append(ent.text)
#         elif ent.label_ == 'WORK_OF_ART':
#             sc_artwork.append(ent.text)
#         elif ent.label_ == 'LANGUAGE':
#             sc_language.append(ent.text)
#         elif ent.label_ == 'PERCENT':
#             sc_percentage.append(ent.text)
#         elif ent.label_ == 'TIME':
#             sc_time.append(ent.text)
#         elif ent.label_ == 'MONEY':
#             sc_money.append(ent.text)
#         elif ent.label_ == 'QUANTITY':
#             sc_quantity.append(ent.text)
#         elif ent.label_ == 'ORDINAL':
#             sc_ordinal.append(ent.text)
#         elif ent.label_ == 'CARDINAL':
#             sc_cardinal.append(ent.text)
#         elif ent.label_ == 'GPE':
#             if ent.text in countries:
#                 sc_country.append(ent.text)
#             elif ent.text in cities:
#                 sc_city.append(ent.text)
#             else:
#                 sc_state.append(ent.text)

#     df_sc_name = pd.DataFrame({"Name": sc_name})
#     df_sc_org = pd.DataFrame({"Organization": sc_org})
#     df_sc_date = pd.DataFrame({"Date": sc_date})
#     df_sc_city = pd.DataFrame({"City": sc_city})
#     df_sc_state = pd.DataFrame({"State": sc_state})
#     df_sc_country = pd.DataFrame({"Country": sc_country})
#     df_sc_time = pd.DataFrame({"Time": sc_time})
#     df_sc_quantity = pd.DataFrame({"Quantity": sc_quantity})
#     df_sc_ordinal = pd.DataFrame({"Ordinal": sc_ordinal})
#     df_sc_cardinal = pd.DataFrame({"Cardinal": sc_cardinal})
#     df_sc_percentage = pd.DataFrame({"Percentage": sc_percentage})
#     df_sc_event = pd.DataFrame({"Event": sc_event})
#     df_sc_language = pd.DataFrame({"Language": sc_language})
#     df_sc_artwork = pd.DataFrame({"Art Work": sc_artwork})
#     df_sc_money = pd.DataFrame({"Money": sc_money})
#     df_sc_loc = pd.DataFrame({"Location": sc_loc})

# def get_email_addresses(text):
#     global df_get_email

#     get_email = []
#     email_set = []
#     remove_email = []
#     r = re.compile(r'[\w\.-]+@[\w\.-]+')

#     string = text
#     email_ori = r.findall(string)
#     email = [email.lower().strip(".") for email in email_ori]

#     df_get_email = pd.DataFrame({"Email": email})

# def read_pdf(file):
#     pdf_text = ""
#     reader = pypdf.PdfReader(BytesIO(file))
#     for page in reader.pages:
#         pdf_text += page.extract_text()
#     return pdf_text

# @app.post("/upload/")
# async def upload_file(file: UploadFile = File(...)):
#     try:
#         file_path = os.path.join(UPLOAD_DIRECTORY, file.filename)
#         with open(file_path, "wb") as buffer:
#             buffer.write(file.file.read())
#         return {"message": "File uploaded successfully"}
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=str(e))

# @app.post("/extract/")
# async def extract_info():
#     file_path = os.path.join(UPLOAD_DIRECTORY, os.listdir(UPLOAD_DIRECTORY)[0])
#     if not os.path.isfile(file_path):
#         raise HTTPException(status_code=400, detail="No file uploaded")

#     try:
#         mime_type, _ = mimetypes.guess_type(file_path)

#         if mime_type == 'application/pdf':
#             with open(file_path, "rb") as file:
#                 text = read_pdf(file.read())
#         else:
#             raise HTTPException(status_code=400, detail="Unsupported file type")

#         cleaned_text = clean_data(text)
#         spacy_ner_extraction(cleaned_text)
#         get_email_addresses(cleaned_text)

#         extracted_data = {
#             "names": df_sc_name["Name"].tolist(),
#             "organizations": df_sc_org["Organization"].tolist(),
#             "dates": df_sc_date["Date"].tolist(),
#             "cities": df_sc_city["City"].tolist(),
#             "states": df_sc_state["State"].tolist(),
#             "countries": df_sc_country["Country"].tolist(),
#             "emails": df_get_email["Email"].tolist(),
#             # Add other extracted data fields here
#         }

#         return {"text": text, "extracted_data": extracted_data}

#     except Exception as e:
#         raise HTTPException(status_code=500, detail=str(e))








# import PyPDF2
# from fastapi import FastAPI, File, UploadFile, HTTPException
# import os
# import mimetypes
# from PyPDF2 import PdfFileReader
# import pandas as pd
# import spacy
# import geonamescache
# import re
# from io import BytesIO
# from fastapi.middleware.cors import CORSMiddleware
# from cleaner import clean_data

# app = FastAPI()

# nlp = spacy.load("en_core_web_sm")
# # Add CORS middleware
# app.add_middleware(
#     CORSMiddleware,
#     allow_origins=["http://localhost:5173"],  # Allow your frontend origin
#     allow_credentials=True,
#     allow_methods=["*"],  # Allow all HTTP methods
#     allow_headers=["*"],  # Allow all headers
# )
# UPLOAD_DIRECTORY = "uploads"

# if not os.path.exists(UPLOAD_DIRECTORY):
#     os.makedirs(UPLOAD_DIRECTORY)

# def gen_dict_extract(var, key):
#     if isinstance(var, dict):
#         for k, v in var.items():
#             if k == key:
#                 yield v
#             if isinstance(v, (dict, list)):
#                 yield from gen_dict_extract(v, key)
#     elif isinstance(var, list):
#         for d in var:
#             yield from gen_dict_extract(d, key)

# def spacy_ner_extraction(text):
#     global df_sc_name, df_sc_org, df_sc_date, df_sc_city, df_sc_state, df_sc_country, df_sc_time, df_sc_quantity, df_sc_ordinal, df_sc_cardinal, df_sc_percentage, df_sc_event, df_sc_language, df_sc_artwork, df_sc_money, df_sc_loc

#     gc = geonamescache.GeonamesCache()
#     countries = gc.get_countries()
#     cities = gc.get_cities()
#     cities = [*gen_dict_extract(cities, 'name')]
#     countries = [*gen_dict_extract(countries, 'name')]

#     sc_name = []
#     sc_org = []
#     sc_date = []
#     sc_time = []
#     sc_city = []
#     sc_state = []
#     sc_country = []
#     sc_loc = []
#     sc_artwork = []
#     sc_language = []
#     sc_money = []
#     sc_quantity = []
#     sc_ordinal = []
#     sc_cardinal = []
#     sc_percentage = []
#     sc_event = []

#     doc = nlp(text)
#     for ent in doc.ents:
#         if ent.label_ == 'PERSON':
#             sc_name.append(ent.text)
#         elif ent.label_ == 'ORG':
#             sc_org.append(ent.text)
#         elif ent.label_ == 'DATE':
#             sc_date.append(ent.text)
#         elif ent.label_ == 'LOC':
#             sc_loc.append(ent.text)
#         elif ent.label_ == 'WORK_OF_ART':
#             sc_artwork.append(ent.text)
#         elif ent.label_ == 'LANGUAGE':
#             sc_language.append(ent.text)
#         elif ent.label_ == 'PERCENT':
#             sc_percentage.append(ent.text)
#         elif ent.label_ == 'TIME':
#             sc_time.append(ent.text)
#         elif ent.label_ == 'MONEY':
#             sc_money.append(ent.text)
#         elif ent.label_ == 'QUANTITY':
#             sc_quantity.append(ent.text)
#         elif ent.label_ == 'ORDINAL':
#             sc_ordinal.append(ent.text)
#         elif ent.label_ == 'CARDINAL':
#             sc_cardinal.append(ent.text)
#         elif ent.label_ == 'GPE':
#             if ent.text in countries:
#                 sc_country.append(ent.text)
#             elif ent.text in cities:
#                 sc_city.append(ent.text)
#             else:
#                 sc_state.append(ent.text)

#     df_sc_name = pd.DataFrame({"Name": sc_name})
#     df_sc_org = pd.DataFrame({"Organization": sc_org})
#     df_sc_date = pd.DataFrame({"Date": sc_date})
#     df_sc_city = pd.DataFrame({"City": sc_city})
#     df_sc_state = pd.DataFrame({"State": sc_state})
#     df_sc_country = pd.DataFrame({"Country": sc_country})
#     df_sc_time = pd.DataFrame({"Time": sc_time})
#     df_sc_quantity = pd.DataFrame({"Quantity": sc_quantity})
#     df_sc_ordinal = pd.DataFrame({"Ordinal": sc_ordinal})
#     df_sc_cardinal = pd.DataFrame({"Cardinal": sc_cardinal})
#     df_sc_percentage = pd.DataFrame({"Percentage": sc_percentage})
#     df_sc_event = pd.DataFrame({"Event": sc_event})
#     df_sc_language = pd.DataFrame({"Language": sc_language})
#     df_sc_artwork = pd.DataFrame({"Art Work": sc_artwork})
#     df_sc_money = pd.DataFrame({"Money": sc_money})
#     df_sc_loc = pd.DataFrame({"Location": sc_loc})

# def get_email_addresses(text):
#     global df_get_email

#     get_email = []
#     email_set = []
#     remove_email = []
#     r = re.compile(r'[\w\.-]+@[\w\.-]+')

#     string = text
#     email_ori = r.findall(string)
#     email = [email.lower().strip(".") for email in email_ori]

#     df_get_email = pd.DataFrame({"Email": email})

# def read_pdf(file):
#     pdf_text = ""
#     reader = PyPDF2.PdfReader(file)
#     for page in reader.pages:
#         pdf_text += page.extract_text()
#     return pdf_text

# file_path = os.path.join('uploads', 'Lab 3[1].pdf')
# print(read_pdf(file_path))


# @app.post("/upload/")
# async def upload_file(file: UploadFile = File(...)):
#     try:
#         file_path = os.path.join(UPLOAD_DIRECTORY, file.filename)
#         with open(file_path, "wb") as buffer:
#             buffer.write(file.file.read())
#         return {"message": "File uploaded successfully"}
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=str(e))

# @app.post("/extract/")
# async def extract_info():
#     file_path = os.path.join(UPLOAD_DIRECTORY, os.listdir(UPLOAD_DIRECTORY)[0])
#     if not os.path.isfile(file_path):
#         raise HTTPException(status_code=400, detail="No file uploaded")

#     try:
#         mime_type, _ = mimetypes.guess_type(file_path)

#         if mime_type == 'application/pdf':
#             with open(file_path, "rb") as file:
#                 text = read_pdf(file.read())
#         elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
#             # Use python-docx to read docx file
#             from docx import Document
#             with open(file_path, "rb") as file:
#                 doc = Document(file)
#                 text = "\n".join([para.text for para in doc.paragraphs])
#         elif mime_type == 'text/plain':
#             with open(file_path, "r", encoding='utf-8', errors='replace') as file:
#                 text = file.read()
#         else:
#             raise HTTPException(status_code=400, detail="Unsupported file type")

#         spacy_ner_extraction(text)
#         get_email_addresses(text)

#         df_output = pd.concat([
#             df_sc_name, df_get_email, df_sc_org, df_sc_date, df_sc_city, df_sc_state, df_sc_country, df_sc_percentage,
#             df_sc_time, df_sc_quantity, df_sc_cardinal, df_sc_event, df_sc_language, df_sc_artwork, df_sc_money, df_sc_loc
#         ], axis=1)
        
#         clean_data_dict = clean_data(df_output.to_dict(orient='records'))

#         return clean_data_dict
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=str(e))
    





# from fastapi import FastAPI, File, UploadFile, HTTPException
# import pandas as pd
# import spacy
# import geonamescache
# import re
# import os
# import mimetypes
# import fitz  # PyMuPDF
# from docx import Document
# from io import BytesIO
# from cleaner import clean_data


# app = FastAPI()


# nlp = spacy.load("en_core_web_sm")


# UPLOAD_DIRECTORY = "uploads"
# if not os.path.exists(UPLOAD_DIRECTORY):
#     os.makedirs(UPLOAD_DIRECTORY)



# def gen_dict_extract(var, key):
#     if isinstance(var, dict):
#         for k, v in var.items():
#             if k == key:
#                 yield v
#             if isinstance(v, (dict, list)):
#                 yield from gen_dict_extract(v, key)
#     elif isinstance(var, list):
#         for d in var:
#             yield from gen_dict_extract(d, key)

# def spacy_ner_extraction(text):
#     global df_sc_name, df_sc_org, df_sc_date, df_sc_city, df_sc_state, df_sc_country, df_sc_time, df_sc_quantity, df_sc_ordinal, df_sc_cardinal, df_sc_percentage, df_sc_event, df_sc_language, df_sc_artwork, df_sc_money, df_sc_loc

#     gc = geonamescache.GeonamesCache()
#     countries = gc.get_countries()
#     cities = gc.get_cities()
#     cities = [*gen_dict_extract(cities, 'name')]
#     countries = [*gen_dict_extract(countries, 'name')]

#     sc_name = []
#     sc_org = []
#     sc_date = []
#     sc_time = []
#     sc_city = []
#     sc_state = []
#     sc_country = []
#     sc_loc = []
#     sc_artwork = []
#     sc_language = []
#     sc_money = []
#     sc_quantity = []
#     sc_ordinal = []
#     sc_cardinal = []
#     sc_percentage = []
#     sc_event = []

#     doc = nlp(text)
#     for ent in doc.ents:
#         if ent.label_ == 'PERSON':
#             sc_name.append(ent.text)
#         elif ent.label_ == 'ORG':
#             sc_org.append(ent.text)
#         elif ent.label_ == 'DATE':
#             sc_date.append(ent.text)
#         elif ent.label_ == 'LOC':
#             sc_loc.append(ent.text)
#         elif ent.label_ == 'WORK_OF_ART':
#             sc_artwork.append(ent.text)
#         elif ent.label_ == 'LANGUAGE':
#             sc_language.append(ent.text)
#         elif ent.label_ == 'PERCENT':
#             sc_percentage.append(ent.text)
#         elif ent.label_ == 'TIME':
#             sc_time.append(ent.text)
#         elif ent.label_ == 'MONEY':
#             sc_money.append(ent.text)
#         elif ent.label_ == 'QUANTITY':
#             sc_quantity.append(ent.text)
#         elif ent.label_ == 'ORDINAL':
#             sc_ordinal.append(ent.text)
#         elif ent.label_ == 'CARDINAL':
#             sc_cardinal.append(ent.text)
#         elif ent.label_ == 'GPE':
#             if ent.text in countries:
#                 sc_country.append(ent.text)
#             elif ent.text in cities:
#                 sc_city.append(ent.text)
#             else:
#                 sc_state.append(ent.text)

#     df_sc_name = pd.DataFrame({"Name": sc_name})
#     df_sc_org = pd.DataFrame({"Organization": sc_org})
#     df_sc_date = pd.DataFrame({"Date": sc_date})
#     df_sc_city = pd.DataFrame({"City": sc_city})
#     df_sc_state = pd.DataFrame({"State": sc_state})
#     df_sc_country = pd.DataFrame({"Country": sc_country})
#     df_sc_time = pd.DataFrame({"Time": sc_time})
#     df_sc_quantity = pd.DataFrame({"Quantity": sc_quantity})
#     df_sc_ordinal = pd.DataFrame({"Ordinal": sc_ordinal})
#     df_sc_cardinal = pd.DataFrame({"Cardinal": sc_cardinal})
#     df_sc_percentage = pd.DataFrame({"Percentage": sc_percentage})
#     df_sc_event = pd.DataFrame({"Event": sc_event})
#     df_sc_language = pd.DataFrame({"Language": sc_language})
#     df_sc_artwork = pd.DataFrame({"Art Work": sc_artwork})
#     df_sc_money = pd.DataFrame({"Money": sc_money})
#     df_sc_loc = pd.DataFrame({"Location": sc_loc})

# def get_email_addresses(text):
#     global df_get_email

#     get_email = []
#     email_set = []
#     remove_email = []
#     r = re.compile(r'[\w\.-]+@[\w\.-]+')

#     string = text
#     email_ori = r.findall(string)
#     email = [email.lower().strip(".") for email in email_ori]

#     df_get_email = pd.DataFrame({"Email": email})

# def read_pdf(file):
#     pdf_text = ""
#     pdf_document = fitz.open(stream=file, filetype="pdf")
#     for page_num in range(len(pdf_document)):
#         page = pdf_document.load_page(page_num)
#         pdf_text += page.get_text()
#     return pdf_text

# def read_docx(file):
#     doc = Document(BytesIO(file))
#     doc_text = ""
#     for para in doc.paragraphs:
#         doc_text += para.text + "\n"
#     return doc_text

# @app.post("/upload/")
# async def upload_file(file: UploadFile = File(...)):
#     try:
#         file_path = os.path.join(UPLOAD_DIRECTORY, file.filename)
#         with open(file_path, "wb") as buffer:
#             buffer.write(file.file.read())
#         return {"message": "File uploaded successfully"}
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=str(e))

# @app.post("/extract/")
# async def extract_info():
#     file_path = os.path.join(UPLOAD_DIRECTORY, os.listdir(UPLOAD_DIRECTORY)[0])
#     if not os.path.isfile(file_path):
#         raise HTTPException(status_code=400, detail="No file uploaded")

#     try:
      
#         mime_type, _ = mimetypes.guess_type(file_path)

#         if mime_type == 'application/pdf':
#             with open(file_path, "rb") as file:
#                 text = read_pdf(file.read())
#         elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
#             with open(file_path, "rb") as file:
#                 text = read_docx(file.read())
#         elif mime_type == 'text/plain':
#             with open(file_path, "r", encoding='utf-8', errors='replace') as file:
#                 text = file.read()
#         else:
#             raise HTTPException(status_code=400, detail="Unsupported file type")

     
#         spacy_ner_extraction(text)
#         get_email_addresses(text)


#         df_output = pd.concat([
#             df_sc_name, df_get_email, df_sc_org, df_sc_date, df_sc_city, df_sc_state, df_sc_country, df_sc_percentage,
#             df_sc_time, df_sc_quantity, df_sc_cardinal, df_sc_event, df_sc_language, df_sc_artwork, df_sc_money, df_sc_loc
#         ], axis=1)
        
#         clean_data_dict = clean_data(df_output.to_dict(orient='records'))

#         return clean_data_dict
#     except Exception as e:
#         raise HTTPException(status_code=500, detail=str(e))





# # from fastapi import FastAPI, File, UploadFile, HTTPException
# # import os
# # import mimetypes

# # from file_handler import read_pdf, read_docx
# # from info_extraction import spacy_ner_extraction, get_email_addresses
# # from cleaner import clean_data

# # app = FastAPI()

# # UPLOAD_DIRECTORY = "uploads"
# # if not os.path.exists(UPLOAD_DIRECTORY):
# #     os.makedirs(UPLOAD_DIRECTORY)

# # @app.post("/upload/")
# # async def upload_file(file: UploadFile = File(...)):
# #     try:
# #         file_path = os.path.join(UPLOAD_DIRECTORY, file.filename)
# #         with open(file_path, "wb") as buffer:
# #             buffer.write(file.file.read())
# #         return {"message": "File uploaded successfully🎊"}
# #     except Exception as e:
# #         raise HTTPException(status_code=500, detail=str(e))

# # @app.post("/extract/")
# # async def extract_info():
# #     file_path = os.path.join(UPLOAD_DIRECTORY, os.listdir(UPLOAD_DIRECTORY)[0])
# #     if not os.path.isfile(file_path):
# #         raise HTTPException(status_code=400, detail="No file uploaded")

# #     try:
# #         mime_type, _ = mimetypes.guess_type(file_path)

# #         if mime_type == 'application/pdf':
# #             with open(file_path, "rb") as file:
# #                 text = read_pdf(file.read())
# #         elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
# #             with open(file_path, "rb") as file:
# #                 text = read_docx(file.read())
# #         elif mime_type == 'text/plain':
# #             with open(file_path, "r", encoding='utf-8', errors='replace') as file:
# #                 text = file.read()
# #         else:
# #             raise HTTPException(status_code=400, detail="Unsupported file type")

# #         spacy_ner_extraction(text)
# #         get_email_addresses(text)

        
# #         df_output = pd.concat([
# #             df_sc_name, df_get_email, df_sc_org, df_sc_date, df_sc_city, df_sc_state, df_sc_country, df_sc_percentage,
# #             df_sc_time, df_sc_quantity, df_sc_cardinal, df_sc_event, df_sc_language, df_sc_artwork, df_sc_money, df_sc_loc
# #         ], axis=1)
        
# #         clean_data_dict = clean_data(df_output.to_dict(orient='records'))

# #         return clean_data_dict
# #     except Exception as e:
# #         raise HTTPException(status_code=500, detail=str(e))
